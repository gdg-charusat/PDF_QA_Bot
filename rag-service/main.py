from fastapi import FastAPI, Request, File, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, Field
from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_core.documents import Document
from dotenv import load_dotenv
from transformers import AutoConfig, AutoTokenizer, AutoModelForSeq2SeqLM, AutoModelForCausalLM
from slowapi import Limiter, _rate_limit_exceeded_handler
from slowapi.util import get_remote_address
from pathlib import Path
import os 
import re
import uuid
import torch
import uvicorn
import torch
import time
import docx
from slowapi import Limiter
from slowapi.util import get_remote_address
import threading
from datetime import datetime

# IMPORTANT: Authentication REMOVED as per issue requirement
# (Authentication was breaking existing endpoints)

load_dotenv()

app = FastAPI(
    title="PDF QA Bot API",
    description="PDF Question-Answering Bot (Session-based, No Auth)",
    version="2.1.0"
)

# CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Rate Limiter
limiter = Limiter(key_func=get_remote_address)
app.state.limiter = limiter
app.add_exception_handler(RateLimitExceeded, _rate_limit_exceeded_handler)

# ===============================
# SESSION STORAGE (REQUIRED: keep sessionId)
# ===============================
# Format: { session_id: { "vectorstores": [FAISS], "last_accessed": float } }
sessions = {}
SESSION_TIMEOUT = 3600  # 1 hour

# ===============================
# QUERY CONDENSATION PROMPT
# ===============================
CONDENSE_QUESTION_PROMPT = """Given the following conversation history and a follow-up question, 
rewrite the follow-up question to be a standalone, self-contained question that can be understood 
without the conversation history. Replace pronouns (it, they, this, that, he, she) with the actual 
entities they refer to from the conversation history.

Conversation History:
{history}

Follow-up Question: {question}

Standalone Question:"""

# ---------------------------------------------------------------------------
# GLOBAL STATE MANAGEMENT (Thread-safe, Multi-user support)
# ---------------------------------------------------------------------------
# Per-user/session storage with proper cleanup and locking
sessions = {}  # {session_id: {"vectorstore": FAISS, "upload_time": datetime}}
sessions_lock = threading.RLock()  # Thread-safe access to sessions

# Generation model globals (lazy-loaded)
generation_model = None
generation_tokenizer = None
generation_is_encoder_decoder = None

# Load local embedding model (unchanged — FAISS retrieval stays the same)
embedding_model = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")

config = AutoConfig.from_pretrained(HF_GENERATION_MODEL)
is_encoder_decoder = bool(getattr(config, "is_encoder_decoder", False))
tokenizer = AutoTokenizer.from_pretrained(HF_GENERATION_MODEL)

if is_encoder_decoder:
    model = AutoModelForSeq2SeqLM.from_pretrained(HF_GENERATION_MODEL)
else:
    model = AutoModelForCausalLM.from_pretrained(HF_GENERATION_MODEL)

if torch.cuda.is_available():
    model = model.to("cuda")

model.eval()

# ===============================
# REQUEST MODELS
# ===============================
class AskRequest(BaseModel):
    question: str = Field(..., min_length=1)
    session_ids: list = []


class SummarizeRequest(BaseModel):
    session_ids: list = []


def normalize_answer(text: str) -> str:
    """
    Post-processes the LLM-generated answer.
    """
    text = normalize_spaced_text(text)
    text = re.sub(r"^(Answer[^:]*:|Context:|Question:)\s*", "", text, flags=re.I)
    return text.strip()


def condense_question(question: str, history: list) -> str:
    """
    Condense a follow-up question using conversation history into a standalone query.
    This handles pronouns and context-dependent references.
    """
    if not history or len(history) < 1:
        return question
    
    # Build history context from last 3-5 messages
    history_text = ""
    for msg in history[-5:]:
        role = msg.get("role", "")
        content = msg.get("content", "")
        if role and content:
            history_text += f"{role}: {content}\n"
    
    if not history_text.strip():
        return question
    
    # Use LLM to rewrite the question
    prompt = CONDENSE_QUESTION_PROMPT.format(
        history=history_text.strip(),
        question=question
    )
    
    try:
        condensed = generate_response(prompt, max_new_tokens=128)
        condensed = condensed.strip()
        
        # Return condensed question if valid, otherwise original
        if condensed and len(condensed) > 3:
            return condensed
    except Exception:
        pass
    
    return question


# ===============================
# DOCUMENT LOADERS
# ===============================
def load_pdf(file_path: str):
    return PyPDFLoader(file_path).load()


def load_txt(file_path: str):
    with open(file_path, "r", encoding="utf-8") as f:
        return [Document(page_content=f.read())]


def load_docx(file_path: str):
    doc = docx.Document(file_path)
    text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    return [Document(page_content=text)]


def load_document(file_path: str):
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        return load_pdf(file_path)
    elif ext == ".docx":
        return load_docx(file_path)
    elif ext in [".txt", ".md"]:
        return load_txt(file_path)
    else:
        raise ValueError("Unsupported file format")


# ===============================
# MODEL LOADING
# ===============================
def load_generation_model():
    global generation_model, generation_tokenizer, generation_is_encoder_decoder

    if generation_model:
        return generation_tokenizer, generation_model, generation_is_encoder_decoder

    config = AutoConfig.from_pretrained(HF_GENERATION_MODEL)
    generation_is_encoder_decoder = bool(config.is_encoder_decoder)

    generation_tokenizer = AutoTokenizer.from_pretrained(HF_GENERATION_MODEL)


# ===============================
# UTILITIES
# ===============================
def cleanup_expired_sessions():
    current_time = time.time()
    expired = [
        sid for sid, data in sessions.items()
        if current_time - data["last_accessed"] > SESSION_TIMEOUT
    ]
    for sid in expired:
        del sessions[sid]


def generate_response(prompt: str, max_new_tokens: int = 200) -> str:
    device = next(model.parameters()).device
    inputs = tokenizer(prompt, return_tensors="pt", truncation=True, max_length=2048)
    inputs = {k: v.to(device) for k, v in inputs.items()}

    output = model.generate(
        **inputs,
        max_new_tokens=max_new_tokens,
        do_sample=False,
        pad_token_id=tokenizer.pad_token_id or tokenizer.eos_token_id,
    )

    if is_encoder_decoder:
        return tokenizer.decode(output[0], skip_special_tokens=True)

    return tokenizer.decode(
        output[0][inputs["input_ids"].shape[1]:],
        skip_special_tokens=True,
    )


# ===============================
# HEALTH ENDPOINTS (kept from enhancement branch)
# ===============================
@app.get("/healthz")
def health_check():
    return {"status": "healthy"}


@app.get("/readyz")
def readiness_check():
    return {"status": "ready"}


# ===============================
# UPLOAD (NO AUTH, RETURNS session_id)
# ===============================
@app.post("/upload")
@limiter.limit("10/15 minutes")
async def upload_file(request: Request, file: UploadFile = File(...)):
    if not file.filename.lower().endswith(".pdf"):
        return {"error": "Only PDF files are supported"}

    session_id = str(uuid4())
    upload_dir = "uploads"
    os.makedirs(upload_dir, exist_ok=True)
    file_path = os.path.join(upload_dir, f"{uuid4().hex}_{file.filename}")

    try:
        with open(file_path, "wb") as buffer:
            buffer.write(await file.read())

        loader = PyPDFLoader(file_path)
        docs = loader.load()

        splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=100
        )
        chunks = splitter.split_documents(docs)

        vectorstore = FAISS.from_documents(chunks, embedding_model)

        sessions[session_id] = {
            "vectorstores": [vectorstore],
            "last_accessed": time.time()
        }

        return {
            "message": "PDF uploaded and processed",
            "session_id": session_id
        }

    except Exception as e:
        return {"error": f"Upload failed: {str(e)}"}


# ===============================
# ASK (USES session_ids — matches fixed App.js)
# ===============================
@app.post("/ask")
@limiter.limit("60/15 minutes")
def ask_question(request: Request, data: AskRequest):
    """
    Answer questions using session-specific PDF context with thread-safe access.
    """
    session_id = request.headers.get("X-Session-ID", "default")
    vectorstore, upload_time = get_session_vectorstore(session_id)
    
    if vectorstore is None:
        return {"answer": "Please upload a PDF first!"}
    
    try:
        # Thread-safe vectorstore access
        with sessions_lock:
            question = data.question
            history = data.history
            conversation_context = ""
            
            if history:
                for msg in history[-5:]:
                    role = msg.get("role", "")
                    content = msg.get("content", "")
                    if role and content:
                        conversation_context += f"{role}: {content}\n"
            
            # QUERY CONDENSATION: Rewrite question to standalone form using history
            standalone_query = condense_question(question, history)
            
            # Search only within current session's vectorstore using condensed query
            docs = vectorstore.similarity_search(standalone_query, k=4)
            if not docs:
                return {"answer": "No relevant context found in the current PDF."}

            context = "\n\n".join([doc.page_content for doc in docs])

            prompt = f"""You are a helpful assistant answering questions ONLY from the provided PDF document.

Conversation History (for context only):
{conversation_context}

Document Context (ONLY reference this):
{context}

Current Question:
{question}

Instructions:
- Answer ONLY using the document context provided above.
- Do NOT use any information from previous documents or conversations outside this context.
- If the answer is not in the document, say so briefly.
- Keep the answer concise (2-3 sentences max).

Answer:"""

            raw_answer = generate_response(prompt, max_new_tokens=512)
            answer = normalize_answer(raw_answer)
            return {"answer": answer}
            
    except Exception as e:
        return {"answer": f"Error processing question: {str(e)}"}

    return {"answer": normalize_answer(answer), "confidence_score": 85}

# ===============================
# SUMMARIZE
# ===============================
@app.post("/summarize")
@limiter.limit("15/15 minutes")
def summarize_pdf(request: Request, data: SummarizeRequest):
    cleanup_expired_sessions()

    if not data.session_ids:
        return {"summary": "No session selected."}

    vectorstores = []
    for sid in data.session_ids:
        session = sessions.get(sid)
        if session:
            vectorstores.extend(session["vectorstores"])

    if not vectorstores:
        return {"summary": "No documents found."}

    docs = []
    for vs in vectorstores:
        docs.extend(vs.similarity_search("Summarize the document", k=6))

    context = "\n\n".join([d.page_content for d in docs])

    prompt = f"Summarize this document:\n\n{context}\n\nSummary:"
    summary = generate_response(prompt, 250)

    return {"summary": summary}


# ===============================
# COMPARE
# ===============================
@app.post("/compare")
@limiter.limit("10/15 minutes")
def compare_documents(request: Request, data: CompareRequest):
    cleanup_expired_sessions()

    if len(data.session_ids) < 2:
        return {"comparison": "Select at least 2 documents."}

    contexts = []
    for sid in data.session_ids:
        session = sessions.get(sid)
        if session:
            vs = session["vectorstores"][0]
            chunks = vs.similarity_search("main topics", k=4)
            text = "\n".join([c.page_content for c in chunks])
            contexts.append(text)

    if len(contexts) < 2:
        return {"comparison": "Not enough documents to compare."}

    combined = "\n\n---\n\n".join(contexts)

    prompt = (
        "Compare the documents below.\n"
        "Give similarities and differences.\n\n"
        f"{combined}\n\nComparison:"
    )

    comparison = generate_response(prompt, 300)
    return {"comparison": comparison}


@app.get("/health")
def health():
    return {"status": "ok"}


if __name__ == "__main__":
    uvicorn.run("main:app", host="0.0.0.0", port=5000)